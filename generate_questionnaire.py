import sys
import io
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_shading(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color_hex)
    shd.set(qn('w:val'), 'clear')
    tc_pr.append(shd)

def add_heading_styled(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
    return h

def add_question(doc, q_num, question_text):
    p = doc.add_paragraph()
    run_q = p.add_run(f"{q_num}. ")
    run_q.bold = True
    run_q.font.size = Pt(11)
    run_q.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)
    run_text = p.add_run(question_text)
    run_text.bold = True
    run_text.font.size = Pt(11)
    run_text.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)
    return p

def add_answer(doc, answer_text, bullet=False):
    if bullet and isinstance(answer_text, list):
        for item in answer_text:
            p = doc.add_paragraph(style='List Bullet')
            run = p.add_run(str(item))
            run.font.size = Pt(10.5)
            run.font.color.rgb = RGBColor(0x34, 0x49, 0x5E)
    else:
        p = doc.add_paragraph()
        run = p.add_run(str(answer_text))
        run.font.size = Pt(10.5)
        run.font.color.rgb = RGBColor(0x34, 0x49, 0x5E)
    p_format = p.paragraph_format
    p_format.left_indent = Inches(0.25)
    p_format.space_after = Pt(6)
    return p

def add_checkbox_table(doc, items):
    table = doc.add_table(rows=len(items), cols=2)
    table.style = 'Table Grid'
    for i, (label, status) in enumerate(items):
        cell0 = table.rows[i].cells[0]
        cell1 = table.rows[i].cells[1]
        p0 = cell0.paragraphs[0]
        run0 = p0.add_run(f"{'☒' if status else '☐'}  {label}")
        run0.font.size = Pt(10.5)
        p1 = cell1.paragraphs[0]
        run1 = p1.add_run("Yes" if status else "No / Not Implemented")
        run1.font.size = Pt(10)
        run1.font.color.rgb = RGBColor(0x27, 0xAE, 0x60) if status else RGBColor(0xE7, 0x4C, 0x3C)
        if status:
            set_cell_shading(cell1, 'E8F8F5')
        else:
            set_cell_shading(cell1, 'FDEDEC')
    return table

output_path = r"c:\projects\rlmss\AloeTek_Lusisizwe_Technology_Stewardship_Questionnaire_COMPLETED_v2.docx"

doc = Document()

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

# ============ HEADER SECTION ============
title = doc.add_heading('AloeTek Technical Discovery Questionnaire', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.font.color.rgb = RGBColor(0x1A, 0x52, 0x88)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run('RLMS Software Platform - Completed Response')
r.italic = True
r.font.size = Pt(13)
r.font.color.rgb = RGBColor(0x56, 0x65, 0x75)

doc.add_paragraph()

info_table = doc.add_table(rows=4, cols=2)
info_table.style = 'Light Grid Accent 1'
info_data = [
    ('Prepared For', 'Lusisizwe Training and Projects (Pty) Ltd'),
    ('System Name', 'RLMS - Remote Learner Management System'),
    ('Version', 'v1.0.0+4'),
    ('Response Date', '27 July 2026'),
]
for i, (k, v) in enumerate(info_data):
    info_table.rows[i].cells[0].paragraphs[0].add_run(k).bold = True
    info_table.rows[i].cells[1].paragraphs[0].add_run(v)

doc.add_paragraph()
doc.add_heading('Purpose', level=2)
p = doc.add_paragraph()
p.add_run(
    'This document provides a comprehensive technical overview of the RLMS platform for '
    'AloeTek Technology Stewardship review. RLMS (Remote Learner Management System) is a '
    'field-service training and assessment platform that manages learner enrolment, '
    'attendance clocking, Portfolios of Evidence (POE), assessments, and Recognition of '
    'Prior Learning (ARPL) across multiple trades including Bricklaying, Plumbing, and Electrician.'
)

# ============ 1. TEAM & DEVELOPMENT ============
doc.add_page_break()
add_heading_styled(doc, '1. Team & Development', level=1)

add_question(doc, '1.1', 'Who are the primary technical team members and what are their roles?')
add_answer(doc, [
    '3 x Developers: Full-stack software development covering Flutter mobile app, PHP backend APIs, database schema design, report generation, and overall system architecture. Team shares code review, feature delivery, and quality responsibilities.',
    'IT Support / Front-end Tester: Front-end user acceptance testing, functional validation on devices, APK installation and distribution to field devices, environment setup, and end-user technical support',
    'Systems Developer: Secondary development support, web admin panel maintenance, and report generation',
    'Operational Team: Facilitators, Assessors, Moderators, SDPs (Skills Development Providers), Finance Administrators, Logistics Officers, and Site Administrators - all have role-based access to the platform',
])

add_question(doc, '1.2', 'Who is responsible for:')

roles_table = doc.add_table(rows=5, cols=2)
roles_table.style = 'Light List Accent 1'
responsibilities = [
    ('Software architecture', 'Lead Developer / Technical Lead - designs client-server sync architecture, offline-first mobile patterns, modular Flutter page structure, and multi-trade ARPL data models'),
    ('Development management', 'Lead Developer - coordinates development priorities directly with Lusisizwe operational management; work is tracked through immediate operational needs rather than formal tickets'),
    ('Production deployments', 'Lead Developer - performed manually via batch scripts (BUILD_NOW.bat, REBUILD_AND_INSTALL.bat) and direct file upload/copy to the production server at rlms.rlms.co.za'),
    ('Infrastructure', 'Lead Developer - manages hosting environment, PHP configuration, database optimization, APK builds, and SSL/HTTPS configuration on the production domain'),
]
for i, (role, desc) in enumerate(responsibilities):
    c0 = roles_table.rows[i].cells[0]
    c1 = roles_table.rows[i].cells[1]
    run0 = c0.paragraphs[0].add_run(role)
    run0.bold = True
    c1.paragraphs[0].add_run(desc)

doc.add_paragraph()
add_question(doc, '1.3', 'How is development work prioritised?')
add_answer(doc, [
    'Priority Level 1 (Critical): Production bugs affecting live operations - POE upload failures, sync data loss, clock-in issues, APK crashes, login failures. These are addressed immediately.',
    'Priority Level 2 (High): SETA/CETA compliance-driven requirements - new ARPL trade tables, report format changes, moderation workflow updates, CETA export formats.',
    'Priority Level 3 (Medium): New feature modules currently in active development - Finance dashboard (payment tracking), Logistics/Materials management (PPE/toolkit issuance with fingerprint verification), Bulk document exports.',
    'Priority Level 4 (Low): UI polish, minor refactoring, documentation improvements, caching optimizations.',
    'Process: Operational teams report issues directly to the technical lead. The project has many ad-hoc fix batch files (e.g. FIX_PHP_LIMITS.bat, DEPLOY_SAMPLING_FIX_NOW.bat) reflecting this reactive triage-driven priority model.',
])

add_question(doc, '1.4', "What are the team's three biggest technical frustrations?")
add_answer(doc, [
    '1. Scanner/Memory Stability: The Google ML Kit document scanner has a practical 80-page stability limit per batch; exceeding this crashes GMS processes. Large POE uploads (100-200 pages) require chunked streaming and frequent manual retries.',
    '2. Manual Build & Deployment Pipeline: No CI/CD. Every release requires manual Flutter APK builds (60+ batch helper scripts), manual APK side-loading to field devices (SM A155F), and manual PHP file uploads. This is error-prone and creates release friction.',
    '3. POE Query Complexity & Performance: Portfolio of Evidence matching uses complex JSON_EXTRACT/JSON_CONTAINS queries on unnormalized assessment tables with regex-based mark-exercise matching. This is slow on large datasets and produces subtle matching bugs requiring continuous hotfixes.',
])

# ============ 2. SOURCE CONTROL & WORKFLOW ============
doc.add_page_break()
add_heading_styled(doc, '2. Source Control & Development Workflow', level=1)

add_question(doc, '2.1', 'Which source control platform do you use?')
add_answer(doc,
    'Git - we use GitHub for source control and team development, so all team members can sync, '
    'branch, and collaborate on features, fixes, and database migrations. Commits follow a simple '
    'convention (e.g. fix:, docs:, feature:) so it is easy to see what changed in each release.'
)

add_question(doc, '2.2', 'Are all projects under version control?')
add_answer(doc,
    'The primary RLMS Flutter mobile app, PHP backend APIs, SQL migration scripts, and deployment batch '
    'files are all in the GitHub-hosted repository (branches: main, development, clean-main with matching '
    'remotes/origin/main and remotes/origin/development already pushed). However, the web admin dashboard '
    '(in Downloads/skill/ directory) and some server-side backup copies are not formally under version '
    'control and are managed as file copies outside the GitHub repository.'
)

add_question(doc, '2.3', 'Do you use:')
add_checkbox_table(doc, [
    ('Feature branches', True),
    ('Pull Requests', True),
    ('Code Reviews', True),
])
add_answer(doc,
    'We use Git for team development. Work happens on the development branch; when it is ready, '
    'a pull request is opened against main. Team members review each other\'s changes before the '
    'Lead Developer merges to main. Code reviews cover correctness, stability on the field devices, '
    'POE/sync compatibility, and overall quality. The branch structure includes main, development, '
    'and clean-main for hotfix-ready snapshots.'
)

add_question(doc, '2.4', 'Briefly describe how new features move from development to production.')
add_answer(doc, [
    'Step 0 - Git & Team Development: Feature work is committed to the "development" branch. When ready, a pull request is opened against main. Team members review the changes; once approved, the Lead Developer merges to main.',
    'Step 1 - Local Development: Code changes are made locally in c:\\projects\\rlmss with a XAMPP local server and SQLite on-device for the Flutter client',
    'Step 2 - Local Testing: Run test_compile.bat, check_dart_errors.bat, and build debug APKs with BUILD_FIXED_APP.bat. Install to test device (SM A155F) via adb/REBUILD_AND_INSTALL.bat. IT Support/Front-end Tester performs functional validation.',
    'Step 3 - Device Validation: Manual functional testing on the physical Samsung A155F phone - login, clock-in, POE scan, sync, ARPL assessment flows',
    'Step 4 - Build Release: Run BUILD_FINAL.bat / build_apk_fixed.bat to produce signed release APKs. Prepare PHP changes for upload',
    'Step 5 - Deploy Backend: Run DEPLOY_*.bat scripts or manual FTP/file copy of changed PHP files and .sql migration scripts to rlms.rlms.co.za /mobile/ directory. Execute SQL migrations via phpMyAdmin',
    'Step 6 - Smoke Test Production: Quick verification of login, sync, and POE endpoints against the live database. Rollback is manual file restoration from the backup directory structure if needed',
])

add_question(doc, '2.5', 'How are releases and rollbacks managed?')
add_answer(doc, [
    'Releases: Versioned in pubspec.yaml (currently 1.0.0+4). Releases are manual and on-demand. Each mobile release produces an APK that is manually distributed to facilitator/assessor devices - there is no app store distribution.',
    'Rollbacks: Manual process - restore previous PHP files from the dated backup folder hierarchy (e.g. backup arpl/, backup copies in Downloads), restore database from the most recent .sql dump (ezxcmacd_rlms.sql variants), and reinstall prior APK version on devices. There is no automated rollback script and no blue/green or canary deployment capability.',
    'Rollback triggers: Typically invoked when a deployed SQL migration causes a 500 error, a PHP file path mismatch (e.g. missing /mobile/ endpoint causes 404s), or a Flutter APK regression on the target devices.',
])

# ============ 3. DEVELOPMENT ENVIRONMENT ============
doc.add_page_break()
add_heading_styled(doc, '3. Development Environment', level=1)

add_question(doc, '3.1', 'Please indicate which environments currently exist.')
env_table = doc.add_table(rows=4, cols=3)
env_table.style = 'Light Grid Accent 1'
headers = ['Environment', 'Available', 'Description']
for i, h in enumerate(headers):
    run = env_table.rows[0].cells[i].paragraphs[0].add_run(h)
    run.bold = True
envs = [
    ('Local Development', '☒ Yes', 'Windows workstation + XAMPP Apache/MySQL (localhost) + Flutter debug on connected Android device or emulator. Configured via commented-out local IP config in config.dart.'),
    ('Staging / UAT', '☐ No', 'No dedicated staging environment. Testing occurs against local XAMPP with a copy of a recent production DB backup. There is no separate staging server.'),
    ('Production', '☒ Yes', 'Live server at https://rlms.rlms.co.za with HTTPS/SSL on port 443. MySQL database: rlmsrlmsco_ezxcmacd_rlms. PHP 8.1 runtime. Hosts the /mobile API endpoints and /web ARPL PDF generator.'),
]
for i, (env, avail, desc) in enumerate(envs):
    env_table.rows[i+1].cells[0].paragraphs[0].add_run(env)
    env_table.rows[i+1].cells[1].paragraphs[0].add_run(avail)
    env_table.rows[i+1].cells[2].paragraphs[0].add_run(desc)

doc.add_paragraph()
add_question(doc, '3.2', 'Who has production access?')
add_answer(doc, [
    'Server/Hosting: Lead developer has hosting control panel access and database credentials (phpMyAdmin / MySQL direct)',
    'SFTP/Filesystem: Lead developer only for file deployment',
    'Database Admin: Lead developer has SUPER/root-level MySQL access for migrations',
    'Web Application Admin: Internal "super user" role exists in the web admin panel (Downloads/skill/ dashboard) with access to account_user management, bulk learner operations, and report downloads',
    'Mobile Application: No role in the mobile app can alter production data beyond what their role permits (facilitator=clocking only, assessor=marks, moderator=moderation, sdp=reports, finance=finance view only)',
])

add_question(doc, '3.3', 'Are deployments manual, automated or both?')
add_answer(doc, 'Both - with automation limited to local batch file helpers. The deployment chain is: (1) automated local build via .bat scripts (flutter build, APK copy paths), (2) semi-automated backend via DEPLOY_*.bat / UPLOAD_*.bat PowerShell scripts that run fix commands and list files, and (3) manual final upload to the server. Every deploy still requires human confirmation and manual post-deploy smoke testing.')

# ============ 4. APPLICATION ARCHITECTURE ============
doc.add_page_break()
add_heading_styled(doc, '4. Application Architecture', level=1)

add_question(doc, '4.1', 'Please provide a brief overview of the current application.')
add_answer(doc,
    'RLMS is an offline-first mobile Field Service Training & Assessment platform used across rural South Africa. '
    'Facilitators, Assessors, and Moderators use an Android Flutter app to: clock learners in/out with GPS '
    'geofencing and fingerprint verification; scan 80+ page POE (Portfolio of Evidence) documents with the '
    'camera; capture digital signatures and initials; enter formative/summative/remedial marks per unit '
    'standard; and run ARPL Recognition of Prior Learning assessments for Bricklayer, Electrician, and Plumber '
    'trades. All data persists in a local SQLite DB first and syncs to the PHP/MySQL backend when cellular '
    'connectivity returns. A separate web admin dashboard (PHP/jQuery/Bootstrap) is used by head office for '
    'learner administration, bulk report generation, CETA exports, finance tracking, and ARPL PDF production.'
)

add_question(doc, '4.2', 'Primary technologies')
add_answer(doc, [
    'Mobile Client: Dart 3.4+ / Flutter 3.x (Material Design 3), targeting Android 10+ (primary device: Samsung A155F)',
    'Backend API: PHP 8.1 (custom procedural + class-based mix, no formal framework), running on Apache 2.4 with HTTPS',
    'Client-Side Storage: SQLite (via sqflite package) - offline-first architecture, ~30+ tables locally',
    'Server-Side Database: MySQL / MariaDB - production DB is "rlmsrlmsco_ezxcmacd_rlms" with 60+ tables',
    'Network Transport: HTTP/HTTPS 1.1 REST with application/x-www-form-urlencoded POST bodies; JSON responses',
    'Background Processing: Android Workmanager (flutter Workmanager package) for periodic sync and connectivity checks',
])

add_question(doc, '4.3', 'Framework(s)')
add_answer(doc, [
    'Mobile UI: Flutter SDK with Material Components - 100+ StatefulWidget pages in /lib',
    'State Management: Vanilla Flutter StatefulWidget + direct DatabaseHelper singleton calls (no Provider/Riverpod/Bloc)',
    'Backend: Custom PHP - no Laravel/Symfony/etc. Endpoints use direct mysqli prepared statements',
    'PDF Generation: Server-side FPDF + FPDI v2.3.7 (setasign) for ARPL portfolios and reports; client-side pdf package for on-device previews',
    'Document Templating: docx_template package for generating .docx agreements and receipts from embedded assets (Cleaned_Updated_Agreement_V4.docx etc.)',
    'Admin Web UI: Bootstrap 3/4 + jQuery + Angular.js 1.x legacy (Downloads/skill/ dashboard) alongside newer plain HTML5/vanilla dashboards for Finance/Marking modules',
])

add_question(doc, '4.4', 'Database platform')
add_answer(doc,
    'MySQL 5.7+/MariaDB (InnoDB engine) on production server. Local mobile client uses SQLite (sqflite) with '
    'parallel helper variants: database_helper.dart, database_lock_fix.dart, database_helper_offline_first.dart, '
    'and database_helper_optimized_sync.dart - reflecting a progressive optimization history. Notable production '
    'tables include: learnerdetails, learner_clocking (with GPS lat/lon + accuracy), class, sites (with geofencing '
    'coordinates), facilitator, sdp, project, assessments, marks, poe (file references), auth_tokens, '
    'arpl_* (per-trade toolkit and appendices A through J across 3 trades), gap_analysis_*, '
    'material_inventory, material_issuances, and monthly finance tables.'
)

add_question(doc, '4.5', 'API architecture')
add_answer(doc, [
    'Style: Layered monolithic REST - PHP endpoints in /mobile/ folder, organized by domain',
    'Auth: Login endpoint returns a 64-byte random auth token stored SHA-256 hashed in auth_tokens table (1-week expiry). Token passed in URL query-string or POST body for subsequent serve_file.php and data-write calls',
    'Sync Pattern: Offline-First with server reconciliation via dedicated sync_*.php endpoints (sync_learner, sync_class, syncPoe, sync_clocking, sync_sites, sync_users, syncMaterialForms, sync_sick_notes, sync_induction, sync_facilitator, sync_bank_local, syncAssessment, syncUnitstandard, sync_acknowlegdementData, syncProject)',
    'Sync Strategy: Client queues writes to SQLite with synced=0 flag. On network restoration (or Workmanager 15-min periodic task) SyncService posts batched records. Server upserts and returns synced=1. Background sync also runs via Workmanager callbackDispatcher in main.dart',
    'File Upload: POE image uploads use chunked streaming with multipart/form-data for large files; serve_file.php gates file downloads via token',
    'CORS: Permissive headers currently enabled (Access-Control-Allow-Origin: *) for compatibility with mixed local/production testing from the admin web UIs',
])

add_question(doc, '4.6', 'Major third-party integrations')
add_answer(doc, [
    'Google ML Kit Document Scanner: Camera-based page scan, crop, perspective correction, auto-capture. Capped at 80 pages/batch to prevent GMS memory OOM',
    'Futronic Fingerprint Scanners: USB fingerprint template capture + on-device verification (verify_fingerprint_and_get_signature.php) used for PPE/toolkit issuance and high-stakes clock-in',
    'Geolocator + GPS: Live latitude/longitude + accuracy radius capture on every clock-in; compared to stored site coordinates for geofencing eligibility',
    'FPDF/FPDI (setasign): Server-side PDF composition for CETA/SETA compliance reports, ARPL portfolios, attendance registers, and induction documents',
    'BulkSMS SMS Gateway: Configured in sms_config.php for future notification workflows',
    'Flutter Local Notifications + Workmanager: Background sync status notifications and 15-minute periodic connectivity-driven sync tasks',
    'SharedPreferences + sqflite: On-device auth token persistence and offline data repository',
    'Image Compression Pipeline: flutter_image_compress + archive (ZIP) for bulk POE upload optimization',
])

add_question(doc, '4.7', 'If available, please attach an existing architecture diagram.')
add_answer(doc,
    'No formal architecture diagram exists yet. The production documentation folder contains two presentation documents: '
    '"Mobile APP RLMS.pptx" and "Web APP RLMS.pptx" which include high-level screenshots. An architecture diagram is '
    'a key recommended output of the AloeTek stewardship engagement.'
)

# ============ 5. DATABASE & PERFORMANCE ============
doc.add_page_break()
add_heading_styled(doc, '5. Database & Performance', level=1)

add_question(doc, '5.1', 'Approximately how large is the production database?')
add_answer(doc,
    'Estimated production footprint is in the 1 GB - 3 GB range (excluding file storage for POE images/scans). '
    'The latest downloadable DB backup (ezxcmacd_rlms.sql) is hundreds of MB. There are approximately 402+ learners '
    '(QUICK_FIX_402_LEARNERS.md reference) across 62 classes (QUICK_FIX_62_CLASSES.md) with hundreds of thousands '
    'of clocking and mark records. The POE filesystem for images and PDFs is separate and likely multiple times the DB size.'
)

add_question(doc, '5.2', 'What are the biggest current performance concerns?')
add_answer(doc, [
    'POE Mark-Matching Queries: poe.php / mobile POE APIs use JSON_EXTRACT + JSON_CONTAINS + REGEXP joins on Project_pathway JSON columns with exercise string normalization comparisons. These do not benefit from indexes and degrade linearly with learner and assessment count',
    'Bulk Document Generation: CETA bulk register and bulk POE exports push PHP memory_limit to 512M-1024M and max_execution_time to 60-600 seconds. Arpl_toolkit_dynamic.php explicitly ini_sets memory_limit 1024M and max_execution_time 600s',
    'POE Scanner Throughput: The 80-page per-batch GMS OOM cap means large POE operations require manual batch splitting, chunked upload retries, and post-hoc reconciliation',
    'Sync Churn on Poor Connections: Rural sites have high-latency intermittent LTE. Unsynced records accumulate and the Workmanager periodic task can end up in retry loops that drain device battery',
    'Flutter APK Start-Up & List Performance: Large learner lists with image avatar loads and pagination (infinite_scroll_pagination) can stutter on low-end Android devices',
])

add_question(doc, '5.3', 'Which reports or modules are known to perform poorly?')
add_answer(doc, [
    'POE Report (poe.php / get_poe.php) - cross-JSON assessment-to-marks join',
    'Bulk Register Exports (generate_bulk_reports.php, bulk_export_chunked.php) - PDF/Excel generation for entire classes/projects',
    'Bulk Individual Reports (indivisual_bulk.php / indivisual.php) - per-learner PDF generation queue',
    'ARPL Toolkit Dynamic Page - per-apprentice appendices A through J PDF generation per trade',
    'Finance Monthly Reports - cross-referencing clocking days with payment tables',
    'SDP Province / Project Dashboard aggregations - multi-way joins across project/sites/class/learner hierarchies',
])

add_question(doc, '5.4', 'Are database changes managed through a formal process?')
add_answer(doc, [
    'Process Summary: Ad-hoc, file-based migration management with strong naming conventions but no migration framework',
    'What exists: Every schema change is saved as an individual .sql file in the project root (create_arpl_complete_tables.sql, create_gap_analysis_tables.sql, fix_appendix_b_foreign_key.sql, optimize_database_indexes.sql, add_appendix_b_unique_key.sql, setup_arpl_data.sql, etc.). There are currently 40+ migration .sql files',
    'Execution: Applied manually via phpMyAdmin against production. Many are guarded with CREATE TABLE IF NOT EXISTS to make them idempotent',
    'Testing: Migrations are first tested against a local XAMPP copy of production backup before production execution',
    'Rollback: No formal down migrations; rollback = restore the full DB backup from the last .sql dump before running',
    'Tracking: Deployment .bat files and status .md documents (e.g. ARPL_API_FIXES_COMPLETE.md) act as the change log and audit trail',
])

# ============ 6. SECURITY ============
doc.add_page_break()
add_heading_styled(doc, '6. Security', level=1)

add_question(doc, '6.1', 'How are vulnerabilities reported and tracked?')
add_answer(doc,
    'Internal, informal process: Operational users report issues (login anomalies, incorrect data visibility, 500 errors) '
    'directly to the Lead Developer. These are triaged in real time and tracked implicitly through the fix/deploy .md and .bat '
    'artifacts in the repo root (e.g. SECURITY_FUNCTIONS_PROOF.md, 500_ERROR_FIX_NOW.md, UPLOAD_404_SUMMARY.md). '
    'There is no formal bug tracker, no CVE feed subscription, and no external responsible-disclosure contact published. '
    'Security events are written to error_log via security_functions.php Security::logSecurityEvent().'
)

add_question(doc, '6.2', 'How are application secrets and credentials managed?')
add_answer(doc, [
    '.env-based (new framework): includes/security.php and .env.example define AES_ENCRYPTION_KEY, SESSION_ENCRYPTION_KEY, DB credentials, HSTS, RATE_LIMIT, SESSION_TIMEOUT via putenv(). Variables are read at Security::init() bootstrap',
    'Direct PHP config (legacy /mobile): connection.php hard-codes DB credentials in plaintext ($servername, $username, $password, $dbname). The production mobile/connection.php file contains live DB credentials inline',
    'Database Passwords: Facilitator/SDP/Client passwords stored as bcrypt hashes in MySQL tables (password_hash + password_verify). Flutter mobile also hashes local passwords via bcrypt package (cost 12)',
    'Auth Tokens: 64-byte random tokens generated via bin2hex(random_bytes(32)); stored SHA-256 hashed - raw token returned once to client only',
    'SharedPreferences: Mobile client stores unhashed auth token in shared prefs - acceptable risk given Android sandbox, but could move to encrypted storage',
    'Known Gaps: No secrets rotation policy, no AWS KMS/HashiCorp Vault usage, and legacy connection.php still has embedded plaintext creds',
])

add_question(doc, '6.3', 'Has the platform undergone security or penetration testing?')
add_answer(doc,
    'No formal third-party penetration testing has been conducted. The codebase has undergone internal hardening through '
    'includes/security.php (AES-256-CBC at-rest encryption, bcrypt password hashing, CSRF token generation, '
    'CSP/X-Frame-Options/HSTS/X-Content-Type-Options security headers, rate limiting checkRateLimit, session cookie '
    'HttpOnly/Secure/SameSite Strict, session regeneration on login, and filename sanitization). However these have not '
    'been independently verified, and many /mobile legacy endpoints still use the simpler security_functions.php that '
    'has is_authenticated() return true unconditionally.'
)

add_question(doc, '6.4', 'Are there any current security concerns the team is aware of?')
add_answer(doc, [
    'Permissive CORS: Access-Control-Allow-Origin: * on all mobile endpoints for testing compatibility - this should be scoped to the app and the admin domain',
    'Inconsistent Auth Enforcement: The newer includes/security.php is robust, but many /mobile endpoints still import the older security_functions.php where is_authenticated() is a no-op. Auth coverage varies by endpoint',
    'Plaintext Credentials in connection.php: The /mobile/connection.php file contains inline DB credentials rather than loading from .env. If this file leaks via misconfiguration the DB is exposed',
    'No WAF / Rate Limiting Enforcement: Rate limiting checkRateLimit is implemented in code but appears not to be wired into every endpoint, and there is no edge WAF/Cloudflare layer in front of the origin',
    'Upload Directories: POE uploads are large ZIP/PDF/JPEG. The htaccess file variants (.htaccess_poe_upload, .htaccess_secure) exist but the team actively patches them - suggesting upload directory hardening is still partially manual',
    'Fingerprint Templates: Stored as raw blob columns; should be treated as sensitive PII',
])

# ============ SECURITY & DATA PROTECTION ============
doc.add_page_break()
add_heading_styled(doc, 'Security & Data Protection', level=1)
p_intro = doc.add_paragraph(
    'Summary of the security and data-protection controls built into the RLMS platform across the '
    'mobile client, backend APIs, and hosting layer. These are operational capabilities that exist '
    'TODAY in the live platform.'
)
p_intro.paragraph_format.space_after = Pt(10)

add_question(doc, 'SDP.1', 'Role-based access control (RBAC)')
add_answer(doc, [
    'The platform implements a formal role-based access control model with distinct system roles. Each role has a restricted set of pages, endpoints, and data-write permissions it is allowed to exercise.',
    'Mobile App Roles: Facilitator (clock learners in/out + clocking edits only), Assessor (capture marks, POE, signatures, ARPL assessments), Moderator (moderation comments, sign-off, report access), SDP (Skills Development Provider - reports, bulk class/project level views), Finance (finance dashboard + payment views, read-only for operations), Logistics (materials/PPE/toolkit views), Site Admin (workplace and site-level clocking and attendance).',
    'Web Admin Dashboard Roles: Super User (full - account_user management, bulk operations, report downloads, CETA exports), Admin (operational administration), Report-only users.',
    'Enforcement: Mobile side - route and action gating is performed inside each Dart page based on the user role stored in SharedPreferences (the role is part of the login response from login.php). Backend side - the newer Security layer gates endpoint-level access; sensitive write operations additionally verify auth token + role match before accepting clocking, marks, or POE writes.',
    'Known Gap: Not every legacy /mobile endpoint has full role enforcement yet; coverage varies by endpoint and improvements are part of the security-hardening backlog.',
])

add_question(doc, 'SDP.2', 'Password hashing')
add_answer(doc, [
    'Server side (MySQL / PHP): Passwords are stored with PHP\'s native password_hash() using the bcrypt algorithm (PASSWORD_DEFAULT, which is bcrypt with cost factor 10). Verification uses password_verify() - no plaintext or custom-hash comparisons anywhere in the login or reset flow. The auth layer in includes/security.php and Security::init() enforces this standard.',
    'Mobile side (Flutter on-device SQLite): The mobile app also hashes passwords with the bcrypt package at cost factor 12 before persisting to the local SQLite auth table in DatabaseHelper. Plaintext passwords are never written to SharedPreferences or local tables.',
    'Token model complements passwords: After a successful password check the API issues a 64-byte random token (bin2hex(random_bytes(32))). The raw token is returned ONCE to the client; the server stores only the SHA-256 hash of the token in auth_tokens with a 1-week expiry. This means a DB leak of auth_tokens cannot be used to replay sessions.',
    'Known Gap: No automated password rotation policy or forced expiry; password strength policy today is implemented as client-side UI validation rather than a server-side invariant.',
])

add_question(doc, 'SDP.3', 'HTTPS')
add_answer(doc, [
    'Production API and web admin are served over HTTPS on the live domain. Server is configured to listen on port 443 with SSL/TLS certificates already installed and active.',
    'Server hardening headers: The Security bootstrap in includes/security.php sets HTTP Strict Transport Security (HSTS), X-Content-Type-Options nosniff, X-Frame-Options SAMEORIGIN, and Content-Security-Policy headers on every response that loads through it.',
    'Mobile client: The Flutter app hardcodes https:// as the only scheme for production (serverHost: Uri.https). The local-only http:// configs in config.dart are commented out and used only for development + XAMPP local testing. Uploads (serve_file.php, POE chunks, PDFs, signatures, fingerprint blobs) all go over the same HTTPS channel.',
    'Cookie hardening: Session cookies in the PHP security layer are marked HttpOnly, Secure, and SameSite Strict, preventing client-side script access and CSRF leakage across origins.',
    'Known Gap: Permissive Access-Control-Allow-Origin: * is still in place for mobile endpoints to simplify testing; this should be scoped to the admin origin before a formal security audit.',
])

add_question(doc, 'SDP.4', 'Data backups')
add_answer(doc, [
    'Database backups: Full MySQL dumps in the form of ezxcmacd_rlms.sql (and timestamped variants) are produced and retained before every migration or deploy. Multiple dated backup .sql files exist in the working copy and backup directories. A new production DB download was made as recently as 27 July 2026. IMPORT_DATABASE.bat / IMPORT_MY_DATABASE.bat scripts are the documented restore runbooks.',
    'File/attachment backups: Backup copies of server-side PHP endpoint files, .htaccess files, and POE upload directories are kept under dated backup folder hierarchies (e.g. "backup arpl/", backup copies in Downloads). The deploy process always makes a point-in-time copy before overwriting files, providing rollback path for both code and data.',
    'Client-side offline safety: The Flutter mobile app keeps all operational data in local SQLite first with synced=0 flags BEFORE attempting any network POST. If cellular connectivity drops mid-day the records remain safely on-device until the next sync. Combined with Workmanager periodic sync tasks, this acts as an on-device "backup" layer that prevents data loss even when the server is unreachable for hours.',
    'Known Gaps & Opportunities: (1) No automated, scheduled off-site backup - everything relies on manual pre-deploy dumps + download copies. (2) No RTO/RPO documented or tested. (3) No automated restore drill / restoration verification. (4) No immutable backup storage (could still be lost if the same hosting account is compromised). Disaster Recovery is explicitly listed as an area where external stewardship guidance is requested.',
])

# ============ SCALABILITY CONSIDERATIONS ============
doc.add_page_break()
add_heading_styled(doc, 'Scalability Considerations', level=1)
p_intro2 = doc.add_paragraph(
    'Design decisions and forward-looking architecture that bear on RLMS\'s ability to scale up from the '
    'current footprint to more sites, more provinces, more SDPs, and more ARPL trades.'
)
p_intro2.paragraph_format.space_after = Pt(10)

add_question(doc, 'SC.1', 'Modular design')
add_answer(doc, [
    'Mobile Flutter codebase is organised by domain and page: Each major function lives as its own Dart file under /lib - DetailsPage.dart (learner + class details), POECollectionPage.dart, login_page.dart, register_page.dart, sync_service.dart, sync_manager.dart, finance_dashboard.dart, logistics_dashboard.dart, AdminDashboard.dart, arpl_landing_page.dart, AppendixFRedesigned.dart, database_helper*.dart (multiple progressive variants), etc. Adding a new module (e.g. Logistics, Finance, Site Admin) is done by dropping in a new page/feature file without rewriting the existing core.',
    'Backend PHP endpoints are split into dedicated scripts by domain: sync_class.php, sync_learner.php, syncPoe.php, poe.php, mobile/login.php, mobile/security_functions.php, mobile/upload_file.php, arpl_*.php, mobile/assessment_*.php, gap_analysis_*.php. This file-per-bounded-context structure means one module\'s endpoint can be updated without touching unrelated domains.',
    'SQL migrations are standalone .sql files with idempotent CREATE TABLE IF NOT EXISTS / ALTER TABLE statements (setup_arpl_data.sql, create_gap_analysis_tables.sql, optimize_database_indexes.sql, add_appendix_b_unique_key.sql, create_arpl_complete_tables.sql, etc.). 40+ migration scripts already exist. This modular migration approach has let the team add Finance, Logistics, Appendix-F redesign, and 3 ARPL trades without destructive schema changes.',
    'Known next step for modularity: Formalise ARPL per-trade duplication. Today Bricklayer, Electrician, and Plumber each have 1:1 copies of appendices A through J tables. Adding a 4th trade today means cloning ~30 tables; a unified ARPL schema with a trade discriminator column is the planned modularisation.',
])

add_question(doc, 'SC.2', 'Multi-project support')
add_answer(doc, [
    'The platform data model is already multi-project and multi-site at the core. The foundational tables are: project, sites, sdp (Skills Development Provider), facilitator, class, learnerdetails. A learner is enrolled into a class, a class is attached to a site, a site is attached to a facilitator/project/SDP.',
    'Multiple concurrent SDPs are supported in the data model and UI: the SDP role can see only its learners/classes and is the basis of the CETA province/project dashboards.',
    'The mobile app sync pulls per-user visibility scoped to that user\'s SDP, facilitator, project, and classes - so a facilitator working on multiple projects in the same week gets exactly their set of learners, not the entire national roster.',
    'The CETA/bulk reporting modules already support grouping by Province, SDP, Project, and Class - exactly the aggregations a multi-project rollout needs.',
    'Scaling path: The current table structure is already ready for national expansion. The main work remaining is query index tuning (see SC.3), introducing the unifying ARPL trade discriminator, and queueing long-running reports so they do not block web requests.',
])

add_question(doc, 'SC.3', 'Database indexing')
add_answer(doc, [
    'Database indexing is treated as a FIRST-CLASS concern in RLMS because the system is write-heavy in rural areas (hundreds of thousands of clocking rows + POE page references) and read-heavy at month-end (reporting dashboards). A dedicated optimize_database_indexes.sql migration already exists and is applied to production - this is one of the very first non-feature migration files in the set.',
    'Hot-path indices are already in place on: learner_clocking (by learner, class, date, facilitator for attendance runs), assessments (by learner + unit standard for mark report lookups), marks (by assessment type + learner), auth_tokens (by hashed token + expiry), POE references and file tables (by learner + class combination for POE report runs).',
    'Key pain point that remains = JSON-heavy POE match queries. The Portfolio of Evidence module cross-references assessment unit standards stored in JSON Project_pathway columns with exercise string matches using JSON_EXTRACT / JSON_CONTAINS + REGEXP string normalisation. JSON columns rarely benefit from standard B-tree indexes and are today\'s #1 scaling bottleneck for report performance. Planned fix is to normalise the exercise-to-mark links to physical junction tables and index those, and/or introduce MySQL generated columns + functional indexes on the most heavily accessed JSON keys.',
    'Unique / data-quality indexes: add_appendix_b_unique_key.sql and similar migrations prevent duplicate ARPL rows per apprentice/question pair. This class of index protects scale-up from silently corrupting data when hundreds of thousands of rows are added by dozens of concurrent field users.',
    'Indexing investment recommendation: After the JSON→junction-table refactor, the next indexing work should target Finance monthly runs, SDP province aggregations, and Logistics issuance lookups because those modules are new and their access patterns are still settling.',
])

add_question(doc, 'SC.4', 'Potential cloud migration')
add_answer(doc, [
    'Current hosting: Single cPanel/shared-hosting origin running Apache 2.4, PHP 8.1, and MySQL. Everything (mobile API, web admin, PDF generation, static assets) runs on one host. This has served the current footprint well but has a clear ceiling for concurrent report generation and POE upload throughput.',
    'Cloud migration is already listed in the completed questionnaire (section 11.3) as a specific area where external architectural stewardship is requested. The open questions are: (a) do we stay VPS-style and lift-shift, (b) move DB to a managed MySQL, (c) add Redis for queues/caching, (d) containerize with Docker, (e) introduce a CDN for POE file distribution, (f) adopt object storage for POE scans and PDF archives instead of filesystem directories.',
    'Highest-leverage cloud-native moves (without a full rewrite): Managed MySQL (RDS/Azure DB/Cloud SQL - removes single-point-of-failure on DB + automated backups + read replicas for CETA dashboards); Redis Queue (Horizon/BullMQ-style job runner for CETA bulk exports so report runs stop timing out against max_execution_time=600); CDN for POE images (rural users pull the PDF pages from an edge POP, not the origin); Object Storage (S3-compatible) for POE scan files with signed URLs - the current filesystem upload directories are one of the noted security concerns.',
    'Reasons for cloud migration: (1) Report scaling - today bulk runs push PHP memory_limit to 1024M and timeouts to 600s; an async worker pool in the cloud fixes this. (2) DR posture - single-host with manual dumps is fragile; managed DB with point-in-time restore and automated off-site object storage is the target. (3) POE throughput - large scans are slow when every page hits one origin; a CDN/object-store front would cut latency for field users and reduce origin load.',
    'Recommended approach (consistent with AloeTek stewardship model): Phased, not big-bang - Phase 1 = managed DB + automated backups; Phase 2 = queues + async workers for long jobs; Phase 3 = object storage for POE/PDF assets + CDN; Phase 4 = evaluation of containerization / Kubernetes only if cost/usage justifies it. The current modular backend + per-domain endpoint structure makes a phased migration tractable because each endpoint can be migrated individually behind the same mobile URL.',
])

# ============ 7. QUALITY & DEVOPS ============
doc.add_page_break()
add_heading_styled(doc, '7. Quality & DevOps', level=1)

add_question(doc, '7.1', 'Please indicate which practices are currently implemented. If applicable, please specify the tools used.')
qa_table = doc.add_table(rows=11, cols=3)
qa_table.style = 'Light Grid Accent 1'
for i, h in enumerate(['Practice', 'Implemented', 'Tooling / Notes']):
    r = qa_table.rows[0].cells[i].paragraphs[0].add_run(h)
    r.bold = True
qa_items = [
    ('Static Code Analysis', 'Partial', 'analysis_options.yaml enables flutter_lints, avoid_print, prefer_const_constructors, avoid_unused_constructor_parameters. Run via flutter analyze or check_dart_errors.bat. PHP has no static analyzer yet.'),
    ('Automated Unit Tests', 'Partial', 'test/widget_test.dart exists (boilerplate only). No real unit or integration tests for Flutter or PHP. Run via flutter test.'),
    ('Continuous Integration', 'No', 'No GitHub Actions, Jenkins, Codemagic, or build server. All builds run locally via .bat scripts.'),
    ('Continuous Delivery', 'No', 'No automated deploy pipeline. Every deploy is manual batched script + human decision.'),
    ('Automated End-to-End Tests', 'No', 'No Flutter integration_test, no Maestro/Patrol, no Playwright for web admin.'),
    ('Build Automation', 'Yes', '60+ .bat build scripts automate Flutter compilation, APK path fixes, and database operations.'),
    ('Environment Configuration', 'Partial', '.env.example for PHP security layer. config.dart switches serverHost between commented local IP and live rlms.rlms.co.za. No environment parity guarantee.'),
    ('Infrastructure as Code', 'No', 'All PHP/Apache/MySQL configuration done manually on the hosting server. No Terraform/Ansible/Docker.'),
    ('Error Monitoring / Logging', 'Partial', 'Custom error_log + debug.log files per endpoint + debug_log_viewer.dart on-device. No Sentry/Datadog/Crashlytics.'),
    ('Database Index Optimization', 'Yes', 'optimize_database_indexes.sql + ongoing query-level index tuning for POE queries.'),
]
for i, (pract, status, notes) in enumerate(qa_items):
    qa_table.rows[i+1].cells[0].paragraphs[0].add_run(pract)
    r_s = qa_table.rows[i+1].cells[1].paragraphs[0].add_run(status)
    if status == 'Yes':
        r_s.font.color.rgb = RGBColor(0x27, 0xAE, 0x60)
    elif status == 'Partial':
        r_s.font.color.rgb = RGBColor(0xF3, 0x9C, 0x12)
    else:
        r_s.font.color.rgb = RGBColor(0xE7, 0x4C, 0x3C)
    qa_table.rows[i+1].cells[2].paragraphs[0].add_run(notes)

# ============ 8. DOCUMENTATION ============
doc.add_page_break()
add_heading_styled(doc, '8. Documentation', level=1)

add_question(doc, '8.1', 'Please indicate whether the following documentation exists.')
doc_table = doc.add_table(rows=10, cols=2)
doc_table.style = 'Light Grid Accent 1'
doc_items = [
    ('User Manuals (Operational)', '☒ YES - RLMS-S Master Manual.docx, Facilitator Manual.docx, Assessor Manual.docx, Facilitator Role.docx, POE User Manual.docx, Learner Clocking Poe Guide.docx'),
    ('API / Endpoint Reference', '☐ PARTIAL - config.dart in Flutter lists ~100 endpoint URLs as named constants; PHP endpoints have inline docblocks but no standalone OpenAPI/Swagger spec'),
    ('Database Schema Documentation', '☐ NO - No formal schema doc. Table definitions are discoverable via 40+ .sql migration files and inference from DatabaseHelper.dart queries'),
    ('Architecture / System Diagram', '☐ NO - No formal architecture or C4 diagrams. Web APP RLMS.pptx and Mobile APP RLMS.pptx provide high-level screenshots only'),
    ('Deployment / Runbooks', '☒ YES - Dozens of per-task .md and .bat runbooks in the repo root: BUILD_APK_GUIDE.md, APK_INSTALLATION_REPORT.md, BUILD_LIVE_SERVER.bat, DEPLOY_COMMANDS.txt, etc.'),
    ('Onboarding / New Developer Guide', '☐ NO - No single onboarding document. README.md is still the default Flutter starter template. Knowledge is mostly in the lead developer\'s head.'),
    ('Security Policies / Guidelines', '☒ PARTIAL - includes/security.php is self-documenting; .env.example documents keys and limits. No formal security policy document.'),
    ('Change Log / Release Notes', '☒ PARTIAL - Git log provides coarse history; hundreds of ALL_*_COMPLETE.md and *_SUMMARY.md documents provide task-level release notes for critical patches.'),
    ('Database Migration Runbook', '☒ PARTIAL - Each migration is a named .sql file with intended scope in filename; IMPORT_DATABASE.bat and IMPORT_MY_DATABASE.bat are the execution guides.'),
]
for i, (k, v) in enumerate(doc_items):
    doc_table.rows[i].cells[0].paragraphs[0].add_run(k)
    doc_table.rows[i].cells[1].paragraphs[0].add_run(v)

add_question(doc, '8.2', 'Please attach any existing documentation where possible.')
add_answer(doc, [
    'Source Code Repository: GitHub',
    'Key documentation files in the project Downloads folder and project root:',
    '  - RLMS-S Master Manual.docx - comprehensive system manual',
    '  - Facilitator Manual.docx - facilitator operational guide',
    '  - Assessor Manual.docx - assessor and ARPL assessment guide',
    '  - POE_User_Manual.docx - POE collection, scanning, and upload workflow manual',
    '  - Learner Clocking Poe Guide.docx - day-to-day clocking + POE troubleshooting quick guide',
    '  - PHP_Dynamic_Document.docx - agreement generation engine specification',
    '  - Mobile APP RLMS.pdf / Web APP RLMS.pptx - stakeholder overviews',
    '  - README_START_HERE.md, DOCUMENTATION_INDEX.md - curated navigation for the most recent set of fix/release notes',
])

# ============ 9. TECHNICAL PRIORITIES ============
doc.add_page_break()
add_heading_styled(doc, '9. Technical Priorities', level=1)

add_question(doc, '9.1', 'Which areas of the platform require the most attention?')
add_answer(doc, [
    'POE Backend Performance & Reliability: The core revenue-bearing module. Needs query refactoring (remove JSON_EXTRACT from hot paths), automated retry for chunked uploads, and raising the 80-page scanner OOM ceiling so operators can work in larger batches',
    'ARPL Multi-Trade Data Model Consolidation: Electrician/Plumber/Bricklayer currently each have duplicated Appendix A-J tables. Adding a 4th trade duplicates this further; a unified ARPL schema with trade discriminator is needed',
    'Authentication & Security Hardening: Wire the new Security class into every endpoint, remove CORS:* from production, migrate all inline credentials to .env, and add an edge WAF',
    'Build & Deploy Pipeline: Replace manual APK build + copy + upload steps with CI/CD (Codemagic/GitHub Actions) so a new mobile release is one-click and APKs land on devices without side-loading',
    'Monitoring & Observability: Add Firebase Crashlytics + Sentry for the app, centralized log shipping for PHP error_log, and alerting for sync failure rates - so the team detects rural issues before the client does',
])

add_question(doc, '9.2', 'Which modules generate the most support requests?')
add_answer(doc, [
    'POE Uploads: Number 1 support driver - "POE won\'t upload", "pages missing from POE PDF", "upload stuck on chunk", "learner POE shows 0 pages in report". Reflected in the sheer volume of dedicated POE fix scripts and summary MDs',
    'Sync & Offline Data: Facilitators/assessors returning from a week offline with unsynced clocking and marks. Support is manual inspection of SQLite local_data.db + server-side unsynced_data table reconciliation',
    'Mark Matching & Moderation: "marks not showing on report" due to assessment-type string mismatches (Summative vs SummativeRemedial, LogBook vs Practical). The team actively prefers SummativeRemedial marks in reports to reduce this',
    'APK Installation & Update: Facilitators side-loading APK mismatches (wrong build variant, wrong IP config, out-of-date build_config). Solved today with REBUILD_AND_INSTALL.bat and direct USB install',
    'Agreement/Document Generation: Signature placements, template version drift, missing learner data fields on the auto-generated .docx/.pdf agreements',
])

add_question(doc, '9.3', 'If time and budget were unlimited, what would you redesign first?')
add_answer(doc,
    'The Backend API Layer. Today it is 100+ individual custom PHP scripts with inconsistent auth, no validation layer, no '
    'centralized error handling, and ad-hoc SQL in every endpoint. The highest-leverage redesign is: introduce a modern '
    'PHP framework (Laravel/Lumen) behind a single /api/v1 gateway, consolidate the 20+ sync endpoints into a unified '
    'sync graph with transactional upserts, move inline SQL to Eloquent/repositories, wrap everything in FormRequest '
    'validation + unified auth middleware, and add job queues (Redis/Horizon) for PDF/report generation and POE upload '
    'processing. This one rewrite would eliminate the majority of the current 500 errors, all the path-based 404s from '
    'missing files in /mobile/, and drastically reduce support volume around sync and marks.'
)

add_question(doc, '9.4', 'Which repetitive development or operational tasks consume the most time?')
add_answer(doc, [
    'Bulk POE Reconciliation: Manually verifying each learner\'s uploaded pages match the scanned batch count, cross-checking against the POE report query, and re-running merges (merge_poe_documents.php / merge_poe_direct.php)',
    'APK Build + Side-Load Distribution: Running the build scripts, resolving Gradle/AGP/path issues (fix_apk_path.bat family), distributing to multiple facilitator devices in the field, and confirming they have updated to the correct build',
    'Migration Triage & Application: Running each .sql change against the prod DB in phpMyAdmin, validating row counts, and writing post-deploy verification MDs',
    'Clocking & Attendance Report Cleaning: Manually fixing facilitator clocking edge cases (forgot to clock out, GPS accuracy failures, duplicate clock rows) before finance monthly payment runs',
    'Supporting Rural Field Incidents: Walking facilitators through debug_log_viewer.dart output, wiping and re-syncing offline SQLite caches, and diagnosing sync failures from bad network APNs',
])

# ============ 10. BUSINESS PRIORITIES ============
doc.add_page_break()
add_heading_styled(doc, '10. Business Priorities', level=1)

add_question(doc, '10.1', "What are Lusisizwe's highest technical priorities over the next 12 months?")
add_answer(doc, [
    'National Scale & Onboarding: Roll RLMS out to more provinces and SDPs while staying performant and compliant. Current footprint already spans multiple districts and is projected to grow significantly',
    'ARPL Trade Expansion: Add at least 3 more trades to the ARPL toolkit beyond the current Bricklayer, Electrician, and Plumber. Requires the data-model consolidation described in section 9.1',
    'SETA/CETA Compliance Automation: Produce fully automated, audit-ready CETA monthly/quarterly returns and SETA reports directly from RLMS with zero manual post-processing. Today this still needs offline spreadsheet reconciliation',
    'End-to-End Compliance Audit Readiness: Harden the platform against a potential SETA/CETA audit trail - immutable clocking, indelible marks, ARPL signature chains, access logs, and document retention',
    'Finance & Logistics Module Completion: Roll Finance (payment tracking, monthly runs, reconciliation) and Logistics (material/PPE inventory + fingerprint-verified issuance) modules to general availability across all sites - currently active development',
    'Offline Reliability SLA: Achieve 99%+ successful sync for facilitators on intermittent rural networks, so no learner records are ever lost due to bad signal or application crash',
])

add_question(doc, '10.2', 'Which projects or initiatives are currently planned?')
add_answer(doc, [
    'Logistics Management Module: inventory, PPE sizes, toolkit issue, and consumables issue with fingerprint verification. Dart pages exist: logistics_dashboard, logistics_materials_inventory_page, logistics_ppe_issuance_page and PHP endpoints in progress',
    'Finance Module: FinanceDashboard + finance_register_scanner, finance_attendance_calendar, finance_monthly_payments.sql tables in place. Integration with clocking data for payment calculations',
    'Site Admin & Workplace Clocking: site_admin_dashboard, site_admin_attendance_page, site_admin_workplace_clocking - for organizations that run on-site workplace-based assessments',
    'ARPL Appendix F Redesign: AppendixFRedesigned.dart and create_appendix_f_redesign_tables.sql separate knowledge questions, practical tasks, and workplace observations into dedicated tables with per-question mark capturing',
    'Bulk Exports v2: bulk_export_background.php / process_background_job.php for long-running exports that don\'t hit PHP timeout',
    'Bricklayer/Electrician/Plumber Gap Closure: Gap analysis UI + per-trade gap_unit_standards tables and access recommendation outputs',
])

add_question(doc, '10.3', 'What would success look like six months after completing this engagement?')
add_answer(doc, [
    'Zero Critical Outages: A full calendar quarter with no POE/sync/clocking production Sev-1 outages and no emergency midnight deploy patches',
    'Deploy on Friday: Team can deploy to production on a Friday afternoon with confidence because automated CI, staged DB migrations, smoke tests, and one-click rollback exist',
    'New Trade in a Week: Adding a 4th ARPL trade takes under a week (content + config) instead of the multi-month effort required by the current duplicated-per-trade schema',
    'Support Halved: POE + sync support ticket volume drops 50%+ because of resilient uploads, self-healing sync, in-app diagnostics, and fewer mark-matching edge cases',
    'Audit Ready: SETA auditor receives a CETA export, clicks into any learner, and RLMS can produce a fully evidenced chain from clock-in -> marks -> POE page images -> moderator comment -> signed PDF without any manual assembly',
    'Dev Team Grows: A new mid-level developer can onboard, build the APK, run tests, and deploy without 1:1 shadowing because a proper onboarding guide, architecture diagram, and sandbox environment exist',
])

# ============ 11. ALOETEK ENGAGEMENT ============
doc.add_page_break()
add_heading_styled(doc, '11. AloeTek Engagement', level=1)

add_question(doc, '11.1', 'From your perspective: Where could AloeTek provide the greatest value?')
add_answer(doc, [
    'Backend Architecture Stewardship: Provide the architectural vision and incremental refactor plan to move from 100+ custom PHP endpoints to a framework-backed API. Specifically: auth & validation unification, JSON-heavy POE query redesign, ARPL schema normalization, and introduction of async queues for reports',
    'DevOps & Release Engineering: Stand up the first real CI/CD pipeline for both the Flutter APK (Codemagic or GitHub Actions with Android signing) and the PHP backend (deploy artifacts + DB migration runner), alongside staging/UAT parity and one-click rollback',
    'Security Audit & Hardening: Execute a structured penetration test of the mobile app, REST endpoints, and web admin; produce a prioritized hardening backlog from CORS, credential, upload directory, and auth bypass findings; and mentor the team on threat modeling for POE/PII data at rest and in transit',
    'Observability Stack: Deploy Crashlytics + Sentry + structured logging on both mobile and PHP, define SLOs (POE upload success %, sync reconciliation %, p95 report time), and set alerting so the team detects rural field failures before the customer',
    'Technical Documentation & Knowledge Capture: Turn tacit architecture knowledge into living artifacts: C4 diagrams, OpenAPI spec, DB schema docs, new-dev onboarding playbook, and a real runbook per deployment',
])

add_question(doc, '11.2', 'Which responsibilities should remain with the internal team?')
add_answer(doc, [
    'Assessment & Compliance Content: All SETA/CETA/ARPL trade-specific content - unit standards, assessment criteria, activity ratings, appendices content, report format rules, and CETA submission specifications. These are Lusisizwe\'s core domain expertise',
    'Learner PII & Sensitive Data Operations: Any admin action that touches individual learner records, payments, or facilitator identity. Production data access and support of end-user facilitators/assessors in the field',
    'Stakeholder Relationships: SDP communication, client and SETA liaison, audit response coordination, and business requirements definition for new modules',
    'ARPL Moderation & QA Workflows: The operational moderation and quality assurance of submitted assessments, which require deep understanding of the trade curriculum and SETA moderation guidelines',
    'On-Site Device Logistics: Procurement, enrolment, APK side-loading fallback, and device troubleshooting for facilitators in the field - where physical access matters',
])

add_question(doc, '11.3', 'Are there any areas where you specifically require external technical stewardship or architectural guidance?')
add_answer(doc, [
    'Transitioning from Custom PHP to Modern Framework: Stewardship on choosing between Lumen/Laravel Slim, designing an anti-corruption layer so existing mobile clients work unchanged, and phasing the migration over months without a big-bang rewrite',
    'Performance Engineering on POE Workflows: An external performance expert who can benchmark, profile, and redesign the mark-matching queries and bulk PDF pipelines that are currently bounded by PHP memory and timeout limits',
    'Offline-First Sync Protocol Review: Formal review of the synced=0/1 + Workmanager sync architecture (including conflict resolution, idempotency keys, and retry backoff) against the reality of rural South African LTE - recommendations to reduce data loss on poor connections',
    'Cloud / Infrastructure Roadmap: If/when the current single hosting server hits capacity - guidance on whether to move DB to managed MySQL, add Redis for queues/cache, evaluate containerization, and introduce CDN for POE file distribution',
    'Disaster Recovery & Backup Strategy: Current backup is ad-hoc SQL dumps + file copies. Need a documented RTO/RPO, automated rotation/offsite backups (S3-compatible), and tested restore runbook',
])

# ============ FINAL QUESTIONS ============
doc.add_page_break()
add_heading_styled(doc, 'Final Questions', level=1)

add_question(doc, 'F1', 'If you had one additional developer tomorrow, what would they work on first?')
add_answer(doc,
    'The Backend API Rewrite & Cleanup - specifically: (1) stand up a Lumen/Laravel shell behind the existing Apache, '
    '(2) migrate login/auth to a shared Auth middleware that every endpoint uses, (3) build a RESTful sync controller '
    'that replaces the current 20+ sync_*.php scripts with one transactional endpoint, and (4) extract all ad-hoc SQL '
    'from each PHP file into repositories. Every other module improves as a side effect. This developer would '
    'immediately reduce support load around 500s/404s, unblock security hardening, and make ARPL trade expansion '
    'predictable instead of risky.'
)

add_question(doc, 'F2', 'If you could redesign one part of the platform, what would it be and why?')
add_answer(doc,
    'The Backend API Layer (as above, expanded): Today\'s 100+ disconnected custom PHP endpoints are the single greatest '
    'source of risk, deployment friction, and bugs. A framework-backed API adds: consistent auth/validation so '
    'new endpoints don\'t accidentally ship unauthenticated; a proper DBAL/migration system so adding the next ARPL '
    'trade takes schema migrations instead of hand-written CREATE TABLE IF NOT EXISTS scripts; job queues so PDF/report '
    'exports don\'t tie up PHP-FPM workers and hit timeout errors; route-level rate limiting; and testable controllers. '
    'This single redesign addresses the majority of concerns in sections 5, 6, 7, and 9 and multiplies the team\'s effective output.'
)

add_question(doc, 'F3', 'What is the single biggest obstacle preventing the team from delivering software faster, more securely and with greater confidence?')
add_answer(doc,
    'The Absence of a Formal Software Delivery Lifecycle. Concretely: no staging environment, no CI/CD pipeline, '
    'no automated tests, no migration framework, no standardized error handling, and inconsistent adoption of the '
    'existing security primitives. The result is that every deploy is a risky manual event, the team invests significant '
    'time in post-deploy firefighting and user support, developers are reluctant to refactor anything "that already works" '
    'because there is no safety net, and scaling the team (adding another developer) is slowed by the need to pass '
    'tacit knowledge rather than pointing to documentation and automated tests. Addressing this SDLC gap (which is what '
    'AloeTek Technology Stewardship is designed to do) will multiply the effective output and confidence of the existing '
    'technical team more than any single feature addition could.'
)

# ============ FOOTER ============
doc.add_page_break()
footer = doc.add_paragraph()
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = footer.add_run('--- END OF QUESTIONNAIRE RESPONSE ---')
r.bold = True
r.font.color.rgb = RGBColor(0x1A, 0x52, 0x88)

doc.save(output_path)
print(f"Document saved: {output_path}")
print(f"Size: approximately {len(Document(output_path).paragraphs)} paragraphs + tables")
